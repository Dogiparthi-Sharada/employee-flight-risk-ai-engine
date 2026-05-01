"""
Build attrisense_ieee.docx — IEEE-style 2-column conference paper
Generated from the markdown source (attrisense_ieee.md), formatted as a proper
Word document with title, author block, abstract, two-column body, headings,
tables, equations (as text/images optional), and references.

Run:
    source /global/gtsna_northeast6/vanama/python_venvs/py311/bin/activate
    python build_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
OUT = HERE / "attrisense_ieee.docx"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_two_column_section(doc):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), "360")  # ~0.25in gutter
    return section


def add_one_column_section(doc):
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")[0]
    cols.set(qn("w:num"), "1")
    return section


def style_normal(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0


def add_heading(doc, text, level=1):
    """Custom IEEE-ish headings:
    level 1 -> 'I. Introduction' style, bold, centered, all caps small.
    level 2 -> italic, left.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    if level == 1:
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.italic = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_body_para(doc, text, justify=True, first_line_indent=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    return p


# ----------------------------------------------------------------------
# Build the document
# ----------------------------------------------------------------------

doc = Document()

# Page setup: US Letter, IEEE-ish margins
section0 = doc.sections[0]
section0.page_width = Inches(8.5)
section0.page_height = Inches(11)
section0.top_margin = Inches(0.75)
section0.bottom_margin = Inches(1.0)
section0.left_margin = Inches(0.625)
section0.right_margin = Inches(0.625)

style_normal(doc)

# ============== TITLE BLOCK (single column) ==============

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
trun = title.add_run(
    "AttriSense: A Multi-Modal Framework for Proactive Employee "
    "Flight-Risk Prediction Combining Imbalanced Tabular Learning, "
    "Retrieval-Augmented Generation, and Natural-Language SQL"
)
trun.bold = True
trun.font.size = Pt(20)
trun.font.name = "Times New Roman"
title.paragraph_format.space_after = Pt(12)

# Author block
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
arun = author.add_run("Sharada Dogiparthi")
arun.bold = True
arun.font.size = Pt(11)
arun.font.name = "Times New Roman"

aff = doc.add_paragraph()
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
afrun = aff.add_run(
    "Master of Science in Business Analytics\n"
    "California State University, East Bay\n"
    "Hayward, CA, USA\n"
    "sdogiparthi@horizon.csueb.edu"
)
afrun.italic = True
afrun.font.size = Pt(10)
afrun.font.name = "Times New Roman"
aff.paragraph_format.space_after = Pt(12)

# ============== SWITCH TO 2-COLUMN ==============
add_two_column_section(doc)

# ============== ABSTRACT ==============
ab_h = doc.add_paragraph()
ab_h_run = ab_h.add_run("Abstract—")
ab_h_run.bold = True
ab_h_run.italic = True
ab_h_run.font.size = Pt(9)
ab_h_run.font.name = "Times New Roman"

abstract_text = (
    "Voluntary employee turnover is one of the costliest, least-instrumented "
    "risks in modern enterprises, with direct replacement cost estimated at "
    "1.5–2.0× annual salary for skilled technical roles. Existing "
    "human-resources analytics tools either (i) report descriptive "
    "statistics after the fact, or (ii) provide opaque predictive scores "
    "that managers do not trust and cannot act on. We present AttriSense, "
    "an end-to-end open-source framework that fuses three complementary "
    "modalities: (1) an imbalanced-classification flight-risk model "
    "(Random Forest with SMOTE), (2) a Retrieval-Augmented Generation "
    "(RAG) layer over qualitative exit-interview corpora using a FAISS "
    "dense index, and (3) a Natural-Language-to-SQL (NL2SQL) agent over "
    "a relational HR data warehouse. Together they produce per-employee "
    "risk scores grounded in both quantitative drivers and qualitative "
    "precedent, and expose the underlying data lake to non-technical HR "
    "users through plain English. On a 5,000-employee synthetic enterprise "
    "corpus calibrated to published HR-attrition distributions, AttriSense "
    "achieves ROC-AUC of 0.91, PR-AUC of 0.74, and reduces analyst-mediated "
    "query latency from a median of 2.3 days to under 12 seconds end-to-end. "
    "We release the full system, dataset generator, evaluation harness, and "
    "a reproducible Streamlit dashboard. Reproducibility note: the "
    "implementation was developed collaboratively with large-language-model "
    "coding assistants; all code, prompts, and design decisions were "
    "human-reviewed and are publicly auditable in the project repository."
)
abp = doc.add_paragraph()
abrun = abp.add_run(abstract_text)
abrun.italic = True
abrun.font.size = Pt(9)
abrun.font.name = "Times New Roman"
abp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
abp.paragraph_format.space_after = Pt(6)

# Keywords
kw = doc.add_paragraph()
kw.paragraph_format.space_after = Pt(8)
kr = kw.add_run("Index Terms—")
kr.italic = True
kr.bold = True
kr.font.size = Pt(9)
kr.font.name = "Times New Roman"
kw.add_run(
    "people analytics, employee attrition, flight risk, "
    "imbalanced classification, SMOTE, retrieval-augmented generation, "
    "FAISS, NL2SQL, large language models, explainable AI, business analytics."
).font.size = Pt(9)

# ============== SECTION I — INTRODUCTION ==============
add_heading(doc, "I. Introduction", level=1)

add_body_para(doc,
    "Voluntary attrition imposes a structural cost on knowledge-intensive "
    "firms that compounds beyond the line item of recruiter fees. Bersin [1] "
    "estimates the total replacement cost of a mid-career engineer at "
    "1.5× to 2.0× annual salary once ramp-up time, lost institutional "
    "knowledge, and team productivity disruption are accounted for. Yet "
    "enterprise HR systems remain overwhelmingly descriptive: they answer "
    "\"who left last quarter?\" rather than \"who is about to leave, and "
    "what should the manager do about it on Monday morning?\""
)

add_body_para(doc, "Three gaps explain this stagnation:")
add_bullet(doc, "The trust gap. Predictive HR models, when deployed, behave as black boxes. Managers refuse to act on a score they cannot explain to the employee in question.")
add_bullet(doc, "The qualitative gap. Most predictive systems consume only structured tabular features (tenure, salary, performance rating) and ignore the rich free-text signal available in exit interviews, 1:1 notes, and engagement surveys.")
add_bullet(doc, "The access gap. HR data lakes are typically gated behind analyst SQL queues with multi-day latency, making real-time managerial use impossible.")

add_body_para(doc,
    "This paper introduces AttriSense, a unified framework that addresses "
    "all three gaps in a single open-source system. Our contributions are: "
    "(a) a reproducible imbalanced-classification baseline using Random "
    "Forest with SMOTE, with full per-bucket precision/recall, calibration, "
    "and PR-AUC reporting; (b) a Retrieval-Augmented Generation pipeline "
    "that grounds quantitative risk scores in semantically retrieved "
    "qualitative exit-interview evidence; (c) a schema-aware NL2SQL agent "
    "built on LangChain [16] and GPT-3.5/4-class LLMs [15], with a 50-question "
    "gold evaluation set; and (d) a unified Streamlit dashboard, released "
    "open-source under MIT license alongside a fully synthetic 5,000-employee "
    "dataset generator and FAISS exit-interview corpus."
)

# ============== SECTION II — RELATED WORK ==============
add_heading(doc, "II. Related Work", level=1)

add_heading(doc, "A. Predictive Attrition Modeling", level=2)
add_body_para(doc,
    "Early work on employee attrition prediction relied on logistic regression "
    "and decision trees over IBM-released benchmarks [2]. Saradhi and "
    "Palshikar [3] demonstrated that ensemble methods outperform single "
    "learners on attrition tasks. Sexton et al. [4] introduced "
    "neural-network-based turnover models. More recent work [5] explores "
    "deep-learning architectures but typically reports only accuracy, "
    "masking the severe class imbalance (~10–15% positive rate in real HR "
    "data) that makes accuracy a misleading metric."
)

add_heading(doc, "B. Imbalanced Classification", level=2)
add_body_para(doc,
    "The Synthetic Minority Over-sampling Technique (SMOTE) [6] remains "
    "the dominant approach for handling HR class imbalance. Recent variants "
    "include Borderline-SMOTE and ADASYN. We adopt vanilla SMOTE for "
    "baseline reproducibility but report calibration and PR-AUC—which are "
    "sensitive to imbalance—in addition to accuracy."
)

add_heading(doc, "C. Retrieval-Augmented Generation", level=2)
add_body_para(doc,
    "Lewis et al. [7] formalized RAG as a paradigm for grounding generative "
    "models in external corpora. Subsequent work has applied RAG to legal, "
    "medical, and customer-support domains, but applications to HR free-text "
    "remain underexplored. We use FAISS [8] for dense retrieval over an "
    "exit-interview embedding index."
)

add_heading(doc, "D. Natural-Language Interfaces to Databases", level=2)
add_body_para(doc,
    "NL2SQL has a long history [9]; modern LLM-based agents [10, 11] achieve "
    "state-of-the-art accuracy on Spider and BIRD benchmarks. We adapt the "
    "LangChain SQL agent pattern with schema injection and self-correction, "
    "evaluating against a hand-curated HR question set."
)

add_heading(doc, "E. Explainability and Fairness in HR AI", level=2)
add_body_para(doc,
    "Recent regulatory frameworks (EU AI Act, NYC Local Law 144 [13]) "
    "classify automated employment decision tools as high-risk, mandating "
    "bias audits and transparency. We treat fairness audit and SHAP [12] "
    "explainability as first-class system components rather than afterthoughts."
)

# ============== SECTION III — DATA ==============
add_heading(doc, "III. Data", level=1)

add_heading(doc, "A. Synthetic Enterprise Corpus", level=2)
add_body_para(doc,
    "Real HR data is fundamentally restricted by privacy regulation and "
    "employment law, and most published academic work uses the small "
    "(n=1,470) IBM HR Analytics dataset [2], which fails to capture "
    "department-level heterogeneity. To enable reproducible research at "
    "enterprise scale without exposing real PII, we built a parametric "
    "synthetic generator producing a 5,000-employee corpus calibrated to "
    "published HR-attrition distributions across:"
)
add_bullet(doc, "demographics (age, tenure, department, location);")
add_bullet(doc, "compensation (base salary, bonus, equity tier);")
add_bullet(doc, "performance (rating, promotion velocity, manager span);")
add_bullet(doc, "engagement signals (survey scores, training completion); and")
add_bullet(doc, "outcome labels (leaver/stayer; if leaver, voluntary/involuntary).")

add_body_para(doc,
    "The generator is parameterized by company size, department mix, "
    "attrition base rate, and seed; it produces both the structured "
    "tabular table and 500 synthetic exit-interview free-text narratives "
    "sampled from a templated reason taxonomy (compensation, "
    "career-growth, manager-relationship, work-life-balance, role-fit)."
)

add_heading(doc, "B. Schema and Storage", level=2)
add_body_para(doc,
    "Tabular data is stored in SQLite for portability; the FAISS index of "
    "exit-interview embeddings is persisted as a flat-index file."
)

# ============== SECTION IV — METHODOLOGY ==============
add_heading(doc, "IV. Methodology", level=1)

add_heading(doc, "A. System Architecture", level=2)
add_body_para(doc,
    "AttriSense consists of four cooperating subsystems: (i) the predictive "
    "engine, (ii) the RAG layer, (iii) the NL2SQL agent, and (iv) the "
    "Streamlit presentation layer. The architecture diagram is included as "
    "Fig. 1 in the supplementary repository."
)

add_heading(doc, "B. Predictive Flight-Risk Engine", level=2)
add_body_para(doc,
    "We train a Random Forest classifier (n_trees = 300, max-depth tuned via "
    "5-fold CV) on the tabular feature matrix. Class imbalance is handled "
    "with SMOTE applied only to the training fold to avoid data leakage. "
    "The model output is the calibrated probability p(leave | x), bucketed "
    "into High (p > 0.75), Medium (0.40 < p ≤ 0.75), and Low (p ≤ 0.40). "
    "Feature attributions are computed per-prediction using SHAP [12] "
    "TreeExplainer; the top-3 contributing features are surfaced to the "
    "manager-facing UI."
)

add_heading(doc, "C. RAG Layer over Exit Interviews", level=2)
add_body_para(doc,
    "Each exit-interview narrative is embedded with the OpenAI "
    "text-embedding-3-small model (1,536-dim) and inserted into a FAISS "
    "IndexFlatIP index. At inference, given a high-risk employee e, we form "
    "a retrieval query from the top-3 SHAP feature names plus e's department "
    "and tenure band; we retrieve the top-k = 5 nearest exit-interview "
    "neighbors. The retrieved snippets are passed to a GPT-4-class LLM with "
    "a constrained prompt to produce a manager-facing 3-bullet retention "
    "rationale."
)

add_heading(doc, "D. NL2SQL Agent", level=2)
add_body_para(doc,
    "We use the LangChain SQL agent pattern. The agent is initialized with "
    "the SQLite schema, a few-shot example bank, and a self-correction "
    "loop: if the generated query raises a SQL error, the error is fed "
    "back into the prompt for one retry. We evaluate against a hand-curated "
    "50-question gold set spanning aggregate queries, joins, and filtered "
    "selections."
)

add_heading(doc, "E. Implementation and Reproducibility", level=2)
add_body_para(doc,
    "The system is implemented in Python 3.11 using scikit-learn [14], "
    "imbalanced-learn, LangChain [16], FAISS [8], and Streamlit. The full "
    "source, prompts, and evaluation harness are released under MIT license. "
    "The codebase was developed collaboratively with LLM coding assistants "
    "(GitHub Copilot, Claude 3.7, GPT-4); the use of AI assistants is "
    "explicitly documented in the project repository's AI_CONTRIBUTIONS.md, "
    "in line with emerging norms for reproducible AI-assisted research."
)

# ============== SECTION V — EVALUATION ==============
add_heading(doc, "V. Evaluation", level=1)

add_heading(doc, "A. Predictive Performance", level=2)
add_body_para(doc, "On a stratified 80/20 train/test split (Table I):")

# Add the perf table
table = doc.add_table(rows=5, cols=5)
table.style = "Light Grid Accent 1"
hdr_cells = table.rows[0].cells
for i, h in enumerate(["Bucket", "Precision", "Recall", "F1", "Support"]):
    p = hdr_cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Times New Roman"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

rows_data = [
    ["Low",    "0.96", "0.98", "0.97", "870"],
    ["Medium", "0.71", "0.62", "0.66", "41"],
    ["High",   "0.83", "0.78", "0.80", "89"],
    ["Macro",  "0.83", "0.79", "0.81", "1,000"],
]
for ridx, rdata in enumerate(rows_data, start=1):
    for cidx, val in enumerate(rdata):
        c = table.rows[ridx].cells[cidx]
        p = c.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
        r.font.name = "Times New Roman"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap.add_run("TABLE I: Predictive Performance on Held-Out Test Set")
cap_run.italic = True
cap_run.font.size = Pt(9)
cap_run.font.name = "Times New Roman"
cap.paragraph_format.space_after = Pt(8)

add_body_para(doc,
    "ROC-AUC: 0.91. PR-AUC: 0.74. Brier score (calibration): 0.061. The "
    "PR-AUC is reported alongside ROC-AUC because PR is more sensitive to "
    "the imbalanced positive class."
)

add_heading(doc, "B. NL2SQL Agent Accuracy", level=2)
add_body_para(doc,
    "On the 50-question gold set, the agent achieves 86% exact-match SQL "
    "execution accuracy and 92% semantic-match (correct answer, "
    "syntactically different SQL)."
)

add_heading(doc, "C. End-to-End Latency", level=2)
add_body_para(doc,
    "Median latency for a complete user query (NL2SQL → DB execution → "
    "LLM-formatted answer) is 11.7 seconds, versus a baseline of 2.3 days "
    "for analyst-mediated queries estimated from a small in-program survey."
)

add_heading(doc, "D. Manager Trust Study (Pilot)", level=2)
add_body_para(doc,
    "A pilot Likert-scale survey (n = 12, MSBA cohort) compared manager "
    "comfort acting on a black-box risk score versus a RAG-grounded "
    "explanation. Mean trust rose from 2.4/5 to 4.1/5 (p < 0.01, Wilcoxon "
    "signed-rank). We position this as preliminary; a full IRB-approved "
    "study with HR practitioners is proposed as future work."
)

# ============== SECTION VI — ETHICS & LIMITATIONS ==============
add_heading(doc, "VI. Ethics, Fairness, and Limitations", level=1)

add_heading(doc, "A. Fairness Audit", level=2)
add_body_para(doc,
    "We compute disparate-impact ratios across simulated gender and age-band "
    "groups using the fairlearn package; on the synthetic data, the "
    "High-bucket positive-prediction rate ratio between protected and "
    "reference groups is 0.94 (within the 0.80 \"four-fifths rule\" "
    "threshold). On real HR data, this audit must be re-run; the methodology, "
    "not the number, is the contribution."
)

add_heading(doc, "B. Synthetic-Data Caveat", level=2)
add_body_para(doc,
    "All performance numbers are reported on a synthetic corpus and may not "
    "transfer to real enterprise data without re-tuning. The synthetic "
    "generator was calibrated to public HR distributions but cannot replicate "
    "firm-specific cultural dynamics."
)

add_heading(doc, "C. Surveillance and Worker Dignity", level=2)
add_body_para(doc,
    "Predictive flight-risk systems carry real risk of being weaponized "
    "against workers (preemptive demotion, exclusion from key projects). "
    "We argue—and the open-source license enforces—that AttriSense outputs "
    "should be visible to the employee in question on request, and should "
    "never be used as the sole input to an adverse employment action."
)

add_heading(doc, "D. Regulatory Alignment", level=2)
add_body_para(doc,
    "The system architecture is designed to satisfy NYC Local Law 144 [13] "
    "bias-audit requirements and EU AI Act high-risk-system documentation "
    "expectations."
)

# ============== SECTION VII — CONCLUSION ==============
add_heading(doc, "VII. Conclusion and Future Work", level=1)

add_body_para(doc,
    "We have presented AttriSense, a unified open-source framework that "
    "combines imbalanced-classification flight-risk prediction, RAG-grounded "
    "qualitative explanation, and NL2SQL data-lake access into a single "
    "workflow consumable by non-technical HR users. On a 5,000-employee "
    "synthetic corpus, the system achieves ROC-AUC of 0.91 and reduces "
    "analyst-query latency by four orders of magnitude. Future work includes: "
    "(i) survival analysis with Cox proportional hazards for time-to-event "
    "prediction; (ii) causal-uplift modeling to recommend the intervention "
    "(raise, transfer, mentorship) that maximally reduces risk; (iii) "
    "federated learning across multi-site enterprises; and (iv) a full "
    "IRB-approved manager-trust field study at scale."
)

# ============== ACKNOWLEDGMENTS ==============
add_heading(doc, "Acknowledgments", level=1)
add_body_para(doc,
    "The author thanks the California State University, East Bay Master of "
    "Science in Business Analytics program for the academic foundation, and "
    "acknowledges the use of GitHub Copilot, Claude 3.7, and GPT-4 as coding "
    "collaborators throughout the implementation. All AI-assisted code was "
    "human-reviewed and is publicly auditable in the project repository."
)

# ============== CODE & DATA AVAILABILITY ==============
add_heading(doc, "Code and Data Availability", level=1)
add_body_para(doc,
    "Source code, synthetic dataset generator, evaluation harness, and the "
    "Streamlit dashboard are publicly available under MIT license at "
    "https://github.com/Dogiparthi-Sharada/employee-flight-risk-ai-engine."
)

# ============== REFERENCES ==============
add_heading(doc, "References", level=1)

refs = [
    "[1] J. Bersin, \"The employee experience platform: A new category arrives,\" Deloitte Insights, 2017.",
    "[2] IBM Watson Analytics, \"HR Analytics Employee Attrition & Performance,\" IBM, 2017.",
    "[3] V. V. Saradhi and G. K. Palshikar, \"Employee churn prediction,\" Expert Systems with Applications, vol. 38, no. 3, pp. 1999–2006, 2011.",
    "[4] R. S. Sexton, S. McMurtrey, J. O. Michalopoulos, and A. M. Smith, \"Employee turnover: A neural network solution,\" Computers & Operations Research, vol. 32, no. 10, 2005.",
    "[5] R. Yedida, R. Reddy, R. Vahi, R. Jana, A. GV, and D. Kulkarni, \"Employee attrition prediction,\" arXiv:1806.10480, 2018.",
    "[6] N. V. Chawla, K. W. Bowyer, L. O. Hall, and W. P. Kegelmeyer, \"SMOTE: Synthetic minority over-sampling technique,\" JAIR, vol. 16, pp. 321–357, 2002.",
    "[7] P. Lewis et al., \"Retrieval-augmented generation for knowledge-intensive NLP tasks,\" NeurIPS, 2020.",
    "[8] J. Johnson, M. Douze, and H. Jégou, \"Billion-scale similarity search with GPUs,\" IEEE Trans. Big Data, vol. 7, no. 3, 2019.",
    "[9] I. Androutsopoulos, G. Ritchie, and P. Thanisch, \"Natural language interfaces to databases—An introduction,\" Natural Language Engineering, vol. 1, no. 1, 1995.",
    "[10] N. Rajkumar, R. Li, and D. Bahdanau, \"Evaluating the text-to-SQL capabilities of large language models,\" arXiv:2204.00498, 2022.",
    "[11] M. Pourreza and D. Rafiei, \"DIN-SQL: Decomposed in-context learning of text-to-SQL with self-correction,\" NeurIPS, 2023.",
    "[12] S. M. Lundberg and S.-I. Lee, \"A unified approach to interpreting model predictions,\" NeurIPS, 2017.",
    "[13] New York City Local Law 144 of 2021, \"Automated employment decision tools,\" effective July 2023.",
    "[14] F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" JMLR, vol. 12, pp. 2825–2830, 2011.",
    "[15] T. Brown et al., \"Language models are few-shot learners,\" NeurIPS, 2020.",
    "[16] H. Chase, \"LangChain,\" GitHub, 2023, https://github.com/langchain-ai/langchain.",
]
for r in refs:
    p = doc.add_paragraph()
    rr = p.add_run(r)
    rr.font.size = Pt(8)
    rr.font.name = "Times New Roman"
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)

doc.save(str(OUT))
print(f"✓ Wrote {OUT}")
